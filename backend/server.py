from fastapi import FastAPI, APIRouter, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import uuid
import json
import asyncio
import base64
import re
import aiohttp
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Ensure exports directory exists
EXPORTS_DIR = ROOT_DIR / "exports"
EXPORTS_DIR.mkdir(exist_ok=True)
IMAGES_DIR = ROOT_DIR / "images"
IMAGES_DIR.mkdir(exist_ok=True)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ====== MODELS ======

class SettingsUpdate(BaseModel):
    api_key_source: str = "emergent"  # "emergent" or "custom"
    custom_api_key: Optional[str] = None
    image_source: str = "ai"  # "ai" or "stock" or "both"
    language: str = "fr"  # "fr" or "en"

class ThemeRequest(BaseModel):
    category: Optional[str] = None
    language: str = "fr"

class IdeaRequest(BaseModel):
    theme: str
    language: str = "fr"

class BookCreateRequest(BaseModel):
    title: str
    subtitle: Optional[str] = None
    description: str
    category: str
    language: str = "fr"
    target_pages: int = 100
    image_source: str = "ai"

class OutlineApproveRequest(BaseModel):
    book_id: str
    outline: List[Dict[str, Any]]

class ChapterGenerateRequest(BaseModel):
    book_id: str

class ExportRequest(BaseModel):
    book_id: str
    format: str = "pdf"  # "pdf", "epub", "docx"

# ====== HELPERS ======

def get_api_key():
    """Get the appropriate API key based on settings."""
    return os.environ.get('EMERGENT_LLM_KEY', '')

async def get_settings():
    settings = await db.settings.find_one({}, {"_id": 0})
    if not settings:
        settings = {
            "api_key_source": "emergent",
            "custom_api_key": None,
            "image_source": "ai",
            "language": "fr"
        }
    return settings

async def get_active_api_key():
    settings = await get_settings()
    if settings.get("api_key_source") == "custom" and settings.get("custom_api_key"):
        return settings["custom_api_key"], "gemini"
    return os.environ.get('EMERGENT_LLM_KEY', ''), "emergent"

async def call_gemini(prompt, system_message="You are a helpful assistant.", session_id=None):
    """Call Gemini 2.5 Flash Lite via emergentintegrations."""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    api_key, key_type = await get_active_api_key()
    if not session_id:
        session_id = str(uuid.uuid4())
    
    chat = LlmChat(
        api_key=api_key,
        session_id=session_id,
        system_message=system_message
    )
    
    if key_type == "gemini":
        # User's own Google key - use directly with Gemini
        chat.with_model("gemini", "gemini-2.5-flash-lite")
    else:
        # Emergent universal key - use gemini-2.5-flash (closest available)
        chat.with_model("gemini", "gemini-2.5-flash")
    
    msg = UserMessage(text=prompt)
    response = await chat.send_message(msg)
    return response

async def generate_image_ai(prompt, book_id, image_name):
    """Generate a photorealistic image using Gemini Nano Banana."""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    api_key, _ = await get_active_api_key()
    
    chat = LlmChat(
        api_key=api_key,
        session_id=str(uuid.uuid4()),
        system_message="You are a professional photographer. Create ultra-realistic, photorealistic images. NEVER create cartoon, illustration, or drawing style images. Always produce images that look like real photographs taken with a professional camera."
    )
    chat.with_model("gemini", "gemini-3-pro-image-preview").with_params(modalities=["image", "text"])
    
    msg = UserMessage(text=prompt)
    text, images = await chat.send_message_multimodal_response(msg)
    
    if images:
        img_data = base64.b64decode(images[0]['data'])
        img_path = IMAGES_DIR / f"{book_id}_{image_name}.png"
        with open(img_path, "wb") as f:
            f.write(img_data)
        return str(img_path), base64.b64encode(img_data).decode('utf-8')[:50]
    return None, None

async def generate_stock_search_query(chapter_title, chapter_content, book_title):
    """Use AI to generate the best stock photo search query for a chapter."""
    try:
        prompt = f"""Based on this book chapter, generate a single short search query (2-4 words) to find a relevant stock photo.
Book: {book_title}
Chapter: {chapter_title}
Content excerpt: {chapter_content[:500]}

Return ONLY the search query, nothing else. Example: "meditation sunrise nature" or "kitchen cooking vegetables" """
        response = await call_gemini(prompt, "Return only a short stock photo search query, no explanation.")
        query = response.strip().strip('"').strip("'")
        # Ensure it's short enough
        words = query.split()[:4]
        return " ".join(words)
    except Exception:
        return chapter_title

async def fetch_stock_image(query, book_id, image_name):
    """Fetch image from free stock sources and save locally."""
    import urllib.parse
    
    encoded_query = urllib.parse.quote(query)
    
    # Use Unsplash Source (no API key needed) - downloads directly
    unsplash_url = f"https://source.unsplash.com/800x600/?{encoded_query}"
    
    async with aiohttp.ClientSession() as session:
        try:
            # Unsplash Source redirects to an actual image
            async with session.get(unsplash_url, timeout=aiohttp.ClientTimeout(total=20), allow_redirects=True) as resp:
                if resp.status == 200 and 'image' in resp.content_type:
                    img_data = await resp.read()
                    if len(img_data) > 1000:  # Ensure we got a real image
                        img_path = IMAGES_DIR / f"{book_id}_{image_name}.png"
                        with open(img_path, "wb") as f:
                            f.write(img_data)
                        return f"/api/images/{book_id}_{image_name}.png"
        except Exception as e:
            logger.error(f"Unsplash fetch error: {e}")
        
        # Fallback: Use Lorem Picsum (always works)
        try:
            picsum_url = f"https://picsum.photos/800/600"
            async with session.get(picsum_url, timeout=aiohttp.ClientTimeout(total=15), allow_redirects=True) as resp:
                if resp.status == 200:
                    img_data = await resp.read()
                    if len(img_data) > 1000:
                        img_path = IMAGES_DIR / f"{book_id}_{image_name}.png"
                        with open(img_path, "wb") as f:
                            f.write(img_data)
                        return f"/api/images/{book_id}_{image_name}.png"
        except Exception as e:
            logger.error(f"Picsum fetch error: {e}")
    
    return None

# ====== SETTINGS ROUTES ======

@api_router.get("/settings")
async def get_settings_route():
    settings = await get_settings()
    if settings.get("custom_api_key"):
        settings["custom_api_key"] = "****" + settings["custom_api_key"][-4:] if len(settings.get("custom_api_key", "")) > 4 else "****"
    return settings

@api_router.put("/settings")
async def update_settings(data: SettingsUpdate):
    update_data = data.model_dump()
    existing = await db.settings.find_one({})
    if existing:
        if update_data.get("custom_api_key") and update_data["custom_api_key"].startswith("****"):
            update_data.pop("custom_api_key")
        await db.settings.update_one({}, {"$set": update_data})
    else:
        await db.settings.insert_one(update_data)
    return {"status": "ok"}

# ====== THEMES ROUTES ======

@api_router.post("/themes/discover")
async def discover_themes(req: ThemeRequest):
    lang = req.language
    category_filter = f" in the category '{req.category}'" if req.category else ""
    
    if lang == "fr":
        prompt = f"""Tu es un expert en analyse de marché Amazon KDP. Analyse les tendances actuelles des livres non-fiction sur Amazon{category_filter}.

Retourne exactement 6 thématiques tendance sous forme de JSON. Chaque thématique doit inclure:
- "title": le nom de la thématique
- "description": une courte description (2 phrases max)
- "demand_level": "high", "medium" ou "low"
- "competition": "high", "medium" ou "low"
- "categories": liste de 2-3 sous-catégories

Concentre-toi sur les livres intemporels: guides pratiques, recettes, tutoriels, développement personnel, livres pour enfants, DIY.

Réponds UNIQUEMENT avec le JSON, sans markdown ni backticks. Format: [{{"title": "...", ...}}]"""
    else:
        prompt = f"""You are an Amazon KDP market analysis expert. Analyze current non-fiction book trends on Amazon{category_filter}.

Return exactly 6 trending themes as JSON. Each theme must include:
- "title": theme name
- "description": short description (2 sentences max)
- "demand_level": "high", "medium" or "low"
- "competition": "high", "medium" or "low"
- "categories": list of 2-3 subcategories

Focus on timeless books: practical guides, recipes, tutorials, self-help, children's books, DIY.

Respond ONLY with JSON, no markdown or backticks. Format: [{{"title": "...", ...}}]"""

    try:
        response = await call_gemini(prompt, "You are an Amazon KDP market expert. Always respond with valid JSON only.")
        # Try to parse JSON from the response
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
            cleaned = cleaned.rsplit("```", 1)[0]
        themes = json.loads(cleaned)
        return {"themes": themes}
    except json.JSONDecodeError:
        logger.error(f"Failed to parse themes JSON: {response[:200]}")
        return {"themes": [], "error": "Failed to parse AI response"}
    except Exception as e:
        logger.error(f"Theme discovery error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ====== IDEAS ROUTES ======

@api_router.post("/ideas/generate")
async def generate_ideas(req: IdeaRequest):
    lang = req.language
    
    if lang == "fr":
        prompt = f"""Tu es un expert en création de livres pour Amazon KDP. Basé sur la thématique "{req.theme}", propose exactement 5 idées de livres non-fiction intemporels.

Pour chaque idée, donne:
- "title": titre accrocheur du livre
- "subtitle": sous-titre descriptif
- "description": description détaillée (3-4 phrases) expliquant le contenu et la valeur ajoutée
- "target_audience": public cible
- "estimated_pages": nombre de pages estimé (entre 80 et 120)
- "category": catégorie (guide, recette, tutoriel, développement personnel, enfants, DIY)
- "unique_angle": ce qui rend ce livre unique

Réponds UNIQUEMENT avec le JSON. Format: [{{"title": "...", ...}}]"""
    else:
        prompt = f"""You are an Amazon KDP book creation expert. Based on the theme "{req.theme}", propose exactly 5 timeless non-fiction book ideas.

For each idea, provide:
- "title": catchy book title
- "subtitle": descriptive subtitle
- "description": detailed description (3-4 sentences) explaining content and value
- "target_audience": target audience
- "estimated_pages": estimated pages (between 80 and 120)
- "category": category (guide, recipe, tutorial, self-help, children, DIY)
- "unique_angle": what makes this book unique

Respond ONLY with JSON. Format: [{{"title": "...", ...}}]"""

    try:
        response = await call_gemini(prompt, "You are a book creation expert. Always respond with valid JSON only.")
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
            cleaned = cleaned.rsplit("```", 1)[0]
        ideas = json.loads(cleaned)
        return {"ideas": ideas}
    except json.JSONDecodeError:
        logger.error(f"Failed to parse ideas JSON: {response[:200]}")
        return {"ideas": [], "error": "Failed to parse AI response"}
    except Exception as e:
        logger.error(f"Ideas generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ====== BOOKS ROUTES ======

@api_router.post("/books/create")
async def create_book(req: BookCreateRequest):
    book_id = str(uuid.uuid4())
    book = {
        "id": book_id,
        "title": req.title,
        "subtitle": req.subtitle,
        "description": req.description,
        "category": req.category,
        "language": req.language,
        "target_pages": req.target_pages,
        "image_source": req.image_source,
        "status": "outline_pending",
        "outline": [],
        "chapters": [],
        "cover_image": None,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.books.insert_one(book)
    book.pop("_id", None)
    return book

@api_router.post("/books/{book_id}/generate-outline")
async def generate_outline(book_id: str):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    lang = book.get("language", "fr")
    pages = book.get("target_pages", 100)
    num_chapters = max(10, pages // 6)
    
    if lang == "fr":
        prompt = f"""Tu es un auteur professionnel. Crée un plan détaillé pour le livre suivant:

Titre: {book['title']}
Sous-titre: {book.get('subtitle', '')}
Description: {book['description']}
Catégorie: {book['category']}
Nombre de pages cible: {pages}

Crée exactement {num_chapters} chapitres. Pour chaque chapitre, donne:
- "chapter_number": numéro du chapitre
- "title": titre du chapitre
- "summary": résumé du contenu (2-3 phrases)
- "key_points": liste de 3-5 points clés à couvrir
- "estimated_pages": pages estimées pour ce chapitre
- "image_suggestion": suggestion d'image pour illustrer ce chapitre

Assure-toi que le total des pages estimées fait environ {pages} pages.
Réponds UNIQUEMENT avec le JSON. Format: [{{"chapter_number": 1, "title": "...", ...}}]"""
    else:
        prompt = f"""You are a professional author. Create a detailed outline for this book:

Title: {book['title']}
Subtitle: {book.get('subtitle', '')}
Description: {book['description']}
Category: {book['category']}
Target pages: {pages}

Create exactly {num_chapters} chapters. For each chapter provide:
- "chapter_number": chapter number
- "title": chapter title
- "summary": content summary (2-3 sentences)
- "key_points": list of 3-5 key points to cover
- "estimated_pages": estimated pages for this chapter
- "image_suggestion": image suggestion to illustrate this chapter

Make sure total estimated pages is around {pages}.
Respond ONLY with JSON. Format: [{{"chapter_number": 1, "title": "...", ...}}]"""

    try:
        response = await call_gemini(prompt, "You are a professional book author. Always respond with valid JSON only.")
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
            cleaned = cleaned.rsplit("```", 1)[0]
        outline = json.loads(cleaned)
        
        await db.books.update_one(
            {"id": book_id},
            {"$set": {"outline": outline, "status": "outline_ready", "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        return {"outline": outline}
    except json.JSONDecodeError:
        logger.error(f"Failed to parse outline JSON: {response[:200]}")
        raise HTTPException(status_code=500, detail="Failed to parse AI response for outline")
    except Exception as e:
        logger.error(f"Outline generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.put("/books/{book_id}/outline")
async def update_outline(book_id: str, req: OutlineApproveRequest):
    await db.books.update_one(
        {"id": book_id},
        {"$set": {"outline": req.outline, "status": "outline_approved", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "ok"}

@api_router.post("/books/{book_id}/generate-chapter/{chapter_num}")
async def generate_chapter(book_id: str, chapter_num: int):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    outline = book.get("outline", [])
    chapter_outline = None
    for ch in outline:
        if ch.get("chapter_number") == chapter_num:
            chapter_outline = ch
            break
    
    if not chapter_outline:
        raise HTTPException(status_code=404, detail="Chapter not found in outline")
    
    lang = book.get("language", "fr")
    est_pages = chapter_outline.get("estimated_pages", 8)
    word_count = est_pages * 250
    
    if lang == "fr":
        prompt = f"""Tu es un auteur professionnel qui écrit le chapitre {chapter_num} du livre "{book['title']}".

Informations du chapitre:
- Titre: {chapter_outline['title']}
- Résumé: {chapter_outline['summary']}
- Points clés: {json.dumps(chapter_outline.get('key_points', []), ensure_ascii=False)}

Écris le contenu complet de ce chapitre. Le texte doit:
- Faire environ {word_count} mots ({est_pages} pages)
- Être professionnel, engageant et bien structuré
- Utiliser ## pour les sous-titres de sections
- Utiliser **texte** pour le gras (mots importants, termes clés)
- Utiliser des listes à puces avec - pour les énumérations
- Couvrir tous les points clés mentionnés
- Inclure des exemples pratiques et des conseils concrets
- Ne PAS utiliser ### ou #### ou *** comme séparateurs
- Ne PAS commencer les paragraphes par des astérisques

Écris UNIQUEMENT le contenu du chapitre, sans meta-commentaires."""
    else:
        prompt = f"""You are a professional author writing chapter {chapter_num} of the book "{book['title']}".

Chapter information:
- Title: {chapter_outline['title']}
- Summary: {chapter_outline['summary']}
- Key points: {json.dumps(chapter_outline.get('key_points', []))}

Write the complete content for this chapter. The text must:
- Be approximately {word_count} words ({est_pages} pages)
- Be professional, engaging and well-structured
- Use ## for section subtitles
- Use **text** for bold (important words, key terms)
- Use bullet lists with - for enumerations
- Cover all key points mentioned
- Include practical examples and concrete advice
- Do NOT use ### or #### or *** as separators
- Do NOT start paragraphs with asterisks

Write ONLY the chapter content, no meta-commentary."""

    try:
        response = await call_gemini(prompt, f"You are writing a professional {book['category']} book. Write detailed, high-quality content.")
        
        chapter_data = {
            "chapter_number": chapter_num,
            "title": chapter_outline["title"],
            "content": response,
            "image_suggestion": chapter_outline.get("image_suggestion", ""),
            "image_url": None,
            "generated_at": datetime.now(timezone.utc).isoformat()
        }
        
        existing_chapters = book.get("chapters", [])
        existing_chapters = [c for c in existing_chapters if c.get("chapter_number") != chapter_num]
        existing_chapters.append(chapter_data)
        existing_chapters.sort(key=lambda x: x.get("chapter_number", 0))
        
        total_chapters = len(outline)
        generated_count = len(existing_chapters)
        new_status = "writing" if generated_count < total_chapters else "chapters_complete"
        
        await db.books.update_one(
            {"id": book_id},
            {"$set": {
                "chapters": existing_chapters,
                "status": new_status,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        return {
            "chapter": chapter_data,
            "progress": {"generated": generated_count, "total": total_chapters}
        }
    except Exception as e:
        logger.error(f"Chapter generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/books/{book_id}/generate-all-chapters")
async def generate_all_chapters_endpoint(book_id: str, background_tasks: BackgroundTasks):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    await db.books.update_one(
        {"id": book_id},
        {"$set": {"status": "writing", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    background_tasks.add_task(generate_all_chapters_task, book_id)
    return {"status": "writing", "message": "Chapter generation started"}

async def generate_all_chapters_task(book_id: str):
    """Background task to generate all chapters one by one."""
    try:
        book = await db.books.find_one({"id": book_id}, {"_id": 0})
        if not book:
            return
        
        outline = book.get("outline", [])
        existing_chapters = book.get("chapters", [])
        generated_nums = {c.get("chapter_number") for c in existing_chapters}
        
        for ch in outline:
            ch_num = ch.get("chapter_number")
            if ch_num in generated_nums:
                continue
            
            try:
                lang = book.get("language", "fr")
                est_pages = ch.get("estimated_pages", 8)
                word_count = est_pages * 250
                
                if lang == "fr":
                    prompt = f"""Tu es un auteur professionnel qui écrit le chapitre {ch_num} du livre "{book['title']}".

Informations du chapitre:
- Titre: {ch['title']}
- Résumé: {ch['summary']}
- Points clés: {json.dumps(ch.get('key_points', []), ensure_ascii=False)}

Écris le contenu complet de ce chapitre ({word_count} mots environ, {est_pages} pages).
Le texte doit être professionnel, engageant, bien structuré.
Utilise ## pour les sous-titres, **texte** pour le gras, - pour les listes.
Ne PAS utiliser ### ou *** comme séparateurs.

Écris UNIQUEMENT le contenu du chapitre."""
                else:
                    prompt = f"""You are a professional author writing chapter {ch_num} of "{book['title']}".

Chapter info:
- Title: {ch['title']}
- Summary: {ch['summary']}
- Key points: {json.dumps(ch.get('key_points', []))}

Write the complete chapter content ({word_count} words, {est_pages} pages).
Be professional, engaging, well-structured.
Use ## for subtitles, **text** for bold, - for lists.
Do NOT use ### or *** as separators.

Write ONLY the chapter content."""

                response = await call_gemini(prompt, f"You are writing a professional {book['category']} book.")
                
                chapter_data = {
                    "chapter_number": ch_num,
                    "title": ch["title"],
                    "content": response,
                    "image_suggestion": ch.get("image_suggestion", ""),
                    "image_url": None,
                    "generated_at": datetime.now(timezone.utc).isoformat()
                }
                
                await db.books.update_one(
                    {"id": book_id},
                    {"$push": {"chapters": chapter_data},
                     "$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
                )
                
            except Exception as e:
                logger.error(f"Error generating chapter {ch_num}: {e}")
                await db.books.update_one(
                    {"id": book_id},
                    {"$set": {"status": "error", "error": f"Chapter {ch_num} failed: {str(e)}"}}
                )
                return
        
        await db.books.update_one(
            {"id": book_id},
            {"$set": {"status": "chapters_complete", "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
    except Exception as e:
        logger.error(f"Background generation error: {e}")
        await db.books.update_one(
            {"id": book_id},
            {"$set": {"status": "error", "error": str(e)}}
        )

@api_router.post("/books/{book_id}/generate-image/{chapter_num}")
async def generate_chapter_image(book_id: str, chapter_num: int):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    chapters = book.get("chapters", [])
    chapter = None
    for c in chapters:
        if c.get("chapter_number") == chapter_num:
            chapter = c
            break
    
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")
    
    settings = await get_settings()
    image_source = settings.get("image_source", "ai")
    ch_title = chapter.get("title", "")
    ch_content = chapter.get("content", "")
    book_title = book.get("title", "")
    
    image_url = None
    if image_source in ("ai", "both"):
        try:
            # Photorealistic AI prompt based on actual chapter content
            content_excerpt = ch_content[:600].replace('\n', ' ')
            prompt = f"""Generate a photorealistic, ultra-realistic photograph for a book chapter.
Book: "{book_title}"
Chapter: "{ch_title}"
Content: {content_excerpt}

The image must look like a real professional photograph, NOT a cartoon, NOT an illustration, NOT a drawing.
High resolution, natural lighting, professional photography style. No text or watermarks in the image."""
            img_path, _ = await generate_image_ai(prompt, book_id, f"ch{chapter_num}")
            if img_path:
                image_url = f"/api/images/{book_id}_ch{chapter_num}.png"
        except Exception as e:
            logger.error(f"AI image generation failed: {e}")
    
    if not image_url and image_source in ("stock", "both"):
        # Use AI to generate a relevant search query from chapter content
        smart_query = await generate_stock_search_query(ch_title, ch_content, book_title)
        logger.info(f"Stock image search query for ch{chapter_num}: '{smart_query}'")
        stock_url = await fetch_stock_image(smart_query, book_id, f"ch{chapter_num}")
        if stock_url:
            image_url = stock_url
    
    if image_url:
        for i, c in enumerate(chapters):
            if c.get("chapter_number") == chapter_num:
                chapters[i]["image_url"] = image_url
                break
        
        await db.books.update_one(
            {"id": book_id},
            {"$set": {"chapters": chapters, "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
    
    return {"image_url": image_url}

@api_router.delete("/books/{book_id}/image/{chapter_num}")
async def delete_chapter_image(book_id: str, chapter_num: int):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    chapters = book.get("chapters", [])
    for i, c in enumerate(chapters):
        if c.get("chapter_number") == chapter_num:
            old_url = c.get("image_url", "")
            chapters[i]["image_url"] = None
            # Delete local file if it exists
            if old_url and old_url.startswith("/api/images/"):
                img_filename = old_url.replace("/api/images/", "")
                img_path = IMAGES_DIR / img_filename
                if img_path.exists():
                    img_path.unlink()
            break
    
    await db.books.update_one(
        {"id": book_id},
        {"$set": {"chapters": chapters, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "deleted"}

@api_router.get("/images/{filename}")
async def serve_image(filename: str):
    file_path = IMAGES_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Image not found")
    return FileResponse(str(file_path), media_type="image/png")

@api_router.get("/books")
async def list_books():
    books = await db.books.find({}, {"_id": 0}).sort("created_at", -1).to_list(100)
    return {"books": books}

@api_router.get("/books/{book_id}")
async def get_book(book_id: str):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    return book

@api_router.delete("/books/{book_id}")
async def delete_book(book_id: str):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    # Delete associated images
    for ch in book.get("chapters", []):
        img_url = ch.get("image_url", "")
        if img_url and img_url.startswith("/api/images/"):
            img_filename = img_url.replace("/api/images/", "")
            img_path = IMAGES_DIR / img_filename
            if img_path.exists():
                try:
                    img_path.unlink()
                except Exception:
                    pass
    
    # Delete export files
    for ext in ["pdf", "docx", "epub"]:
        export_path = EXPORTS_DIR / f"{book_id}.{ext}"
        if export_path.exists():
            try:
                export_path.unlink()
            except Exception:
                pass
    
    result = await db.books.delete_one({"id": book_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Book not found")
    return {"status": "deleted"}

@api_router.get("/books/{book_id}/progress")
async def get_book_progress(book_id: str):
    book = await db.books.find_one({"id": book_id}, {"_id": 0, "chapters.content": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    outline = book.get("outline", [])
    chapters = book.get("chapters", [])
    
    return {
        "status": book.get("status"),
        "total_chapters": len(outline),
        "generated_chapters": len(chapters),
        "chapter_titles": [{"number": c.get("chapter_number"), "title": c.get("title"), "has_image": bool(c.get("image_url"))} for c in chapters],
        "error": book.get("error")
    }

# ====== MARKDOWN HELPERS ======

def parse_markdown_line(line):
    """Parse a markdown line and return (type, content, level)."""
    stripped = line.strip()
    if not stripped:
        return ("blank", "", 0)
    header_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
    if header_match:
        return ("heading", header_match.group(2).strip(), len(header_match.group(1)))
    list_match = re.match(r'^[-*]\s+(.+)$', stripped)
    if list_match:
        return ("list_item", list_match.group(1).strip(), 0)
    num_match = re.match(r'^\d+[.)]\s+(.+)$', stripped)
    if num_match:
        return ("num_list_item", num_match.group(1).strip(), 0)
    if re.match(r'^[-*_]{3,}$', stripped):
        return ("hr", "", 0)
    return ("paragraph", stripped, 0)

def md_to_xml(text):
    """Convert inline markdown to ReportLab XML (bold, italic)."""
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    text = re.sub(r'\*{3}(.+?)\*{3}', r'<b><i>\1</i></b>', text)
    text = re.sub(r'_{3}(.+?)_{3}', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\*{2}(.+?)\*{2}', r'<b>\1</b>', text)
    text = re.sub(r'_{2}(.+?)_{2}', r'<b>\1</b>', text)
    text = re.sub(r'(?<![*])\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)
    text = re.sub(r'`(.+?)`', r'<font face="Courier" size="9">\1</font>', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'<u>\1</u>', text)
    return text

def md_to_html(text):
    """Convert inline markdown to HTML (bold, italic, links)."""
    import html as html_mod
    text = html_mod.escape(text)
    text = re.sub(r'\*{3}(.+?)\*{3}', r'<strong><em>\1</em></strong>', text)
    text = re.sub(r'_{3}(.+?)_{3}', r'<strong><em>\1</em></strong>', text)
    text = re.sub(r'\*{2}(.+?)\*{2}', r'<strong>\1</strong>', text)
    text = re.sub(r'_{2}(.+?)_{2}', r'<strong>\1</strong>', text)
    text = re.sub(r'(?<![*])\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<em>\1</em>', text)
    text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)
    return text

def md_clean(text):
    """Strip all markdown to plain text."""
    text = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.+?)_{1,3}', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def strip_chapter_title_from_content(content, chapter_title):
    """Remove the first heading if it matches or is similar to the chapter title."""
    lines = content.split('\n')
    clean_title = md_clean(chapter_title).strip().lower()
    
    # Find and remove leading blank lines, then check first heading
    start_idx = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            start_idx = i + 1
            continue
        # Check if it's a heading matching the chapter title
        header_match = re.match(r'^#{1,4}\s+(.+)$', stripped)
        if header_match:
            heading_text = md_clean(header_match.group(1)).strip().lower()
            # Check similarity (contains or is very similar)
            if heading_text == clean_title or clean_title in heading_text or heading_text in clean_title:
                start_idx = i + 1
                # Also skip blank line after removed heading
                if start_idx < len(lines) and not lines[start_idx].strip():
                    start_idx += 1
        break
    
    return '\n'.join(lines[start_idx:])

# ====== EXPORT ROUTES ======

@api_router.post("/books/{book_id}/export")
async def export_book(book_id: str, req: ExportRequest):
    book = await db.books.find_one({"id": book_id}, {"_id": 0})
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    
    if not book.get("chapters"):
        raise HTTPException(status_code=400, detail="Book has no chapters")
    
    fmt = req.format.lower()
    
    try:
        if fmt == "pdf":
            filepath = await export_pdf(book)
        elif fmt == "docx":
            filepath = await export_docx(book)
        elif fmt == "epub":
            filepath = await export_epub(book)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")
        
        filename = f"{book['title'].replace(' ', '_')}_{book['id'][:8]}.{fmt}"
        return FileResponse(
            str(filepath),
            media_type="application/octet-stream",
            filename=filename
        )
    except Exception as e:
        logger.error(f"Export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def export_pdf(book):
    """Generate KDP-compliant PDF with accurate page numbers and TOC."""
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, 
        PageBreak, Image, Table, TableStyle, Flowable
    )
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib import colors
    from reportlab.platypus.flowables import HRFlowable
    from io import BytesIO
    
    filepath = EXPORTS_DIR / f"{book['id']}.pdf"
    page_w = 5.5 * inch
    page_h = 8.5 * inch
    left_m = 0.75 * inch
    right_m = 0.5 * inch
    top_m = 0.75 * inch
    bottom_m = 0.75 * inch
    content_w = page_w - left_m - right_m
    
    is_fr = book.get('language') != 'en'
    chapters = sorted(book.get('chapters', []), key=lambda x: x.get('chapter_number', 0))
    
    # ---- Styles ----
    def make_styles():
        s = getSampleStyleSheet()
        s.add(ParagraphStyle('BookTitle', fontName='Times-Bold', fontSize=28,
            alignment=TA_CENTER, leading=34, spaceAfter=12))
        s.add(ParagraphStyle('BookSubtitle', fontName='Times-Italic', fontSize=14,
            alignment=TA_CENTER, textColor=colors.Color(0.4, 0.4, 0.4), spaceAfter=20))
        s.add(ParagraphStyle('ChapLabel', fontName='Helvetica', fontSize=11,
            textColor=colors.Color(0.45, 0.45, 0.45), alignment=TA_CENTER, spaceAfter=8))
        s.add(ParagraphStyle('ChapTitlePage', fontName='Times-Bold', fontSize=22,
            alignment=TA_CENTER, leading=28))
        s.add(ParagraphStyle('H2', fontName='Times-Bold', fontSize=14,
            spaceBefore=16, spaceAfter=8, leading=18))
        s.add(ParagraphStyle('H3', fontName='Times-BoldItalic', fontSize=12,
            spaceBefore=12, spaceAfter=6, leading=16))
        s.add(ParagraphStyle('H4', fontName='Times-Bold', fontSize=11,
            spaceBefore=10, spaceAfter=4, leading=15))
        s.add(ParagraphStyle('Body', fontName='Times-Roman', fontSize=11,
            leading=16, alignment=TA_JUSTIFY, spaceAfter=6))
        s.add(ParagraphStyle('ListItem', fontName='Times-Roman', fontSize=11,
            leading=16, spaceAfter=3, leftIndent=24, bulletIndent=12))
        s.add(ParagraphStyle('TOCTitle', fontName='Times-Bold', fontSize=20,
            spaceAfter=30, alignment=TA_CENTER))
        s.add(ParagraphStyle('TOCLeft', fontName='Times-Roman', fontSize=11, leading=20))
        s.add(ParagraphStyle('TOCRight', fontName='Times-Roman', fontSize=11,
            leading=20, alignment=TA_RIGHT))
        return s
    
    # ---- Chapter marker flowable ----
    class ChapterMark(Flowable):
        """Invisible flowable that records which page a chapter starts on."""
        width = 0
        height = 0
        def __init__(self, ch_num, tracker_dict):
            Flowable.__init__(self)
            self.ch_num = ch_num
            self.tracker_dict = tracker_dict
        def draw(self):
            self.tracker_dict[self.ch_num] = self.canv.getPageNumber()
        def wrap(self, aW, aH):
            return (0, 0)
    
    # ---- Build chapter content flowables ----
    def build_chapter_body(chapter, styles):
        flowables = []
        content = strip_chapter_title_from_content(
            chapter.get('content', ''), chapter.get('title', ''))
        if chapter.get('image_url') and chapter['image_url'].startswith('/api/images/'):
            img_filename = chapter['image_url'].replace('/api/images/', '')
            img_path = IMAGES_DIR / img_filename
            if img_path.exists():
                try:
                    img = Image(str(img_path), width=3.2 * inch, height=2.2 * inch)
                    img.hAlign = 'CENTER'
                    flowables.append(img)
                    flowables.append(Spacer(1, 14))
                except Exception:
                    pass
        for line in content.split('\n'):
            lt, lc, lv = parse_markdown_line(line)
            if lt == "blank":
                flowables.append(Spacer(1, 4))
            elif lt == "heading":
                xml = md_to_xml(lc)
                flowables.append(Paragraph(xml, styles[{1:'H2',2:'H2',3:'H3'}.get(lv,'H4')]))
            elif lt == "list_item":
                flowables.append(Paragraph(f"\u2022  {md_to_xml(lc)}", styles['ListItem']))
            elif lt == "num_list_item":
                flowables.append(Paragraph(f"\u2013  {md_to_xml(lc)}", styles['ListItem']))
            elif lt == "hr":
                flowables.append(Spacer(1, 6))
                flowables.append(HRFlowable(width="60%", thickness=0.5, color=colors.Color(.7,.7,.7)))
                flowables.append(Spacer(1, 6))
            elif lt == "paragraph":
                flowables.append(Paragraph(md_to_xml(lc), styles['Body']))
        return flowables
    
    # ---- Build full story (without real TOC page numbers) ----
    def build_story(styles, toc_page_map=None):
        story = []
        # Title page
        story.append(Spacer(1, 2.5 * inch))
        story.append(Paragraph(md_to_xml(book['title']), styles['BookTitle']))
        if book.get('subtitle'):
            story.append(Spacer(1, 8))
            story.append(Paragraph(md_to_xml(book['subtitle']), styles['BookSubtitle']))
        story.append(PageBreak())
        
        # TOC page
        toc_label = "Table des matieres" if is_fr else "Table of Contents"
        story.append(Paragraph(toc_label, styles['TOCTitle']))
        toc_rows = []
        for ch in chapters:
            ch_num = ch['chapter_number']
            ch_lbl = f"Chapitre {ch_num}" if is_fr else f"Chapter {ch_num}"
            title_text = f"{ch_lbl}  -  {md_to_xml(ch['title'])}"
            page_str = str(toc_page_map.get(ch_num, "")) if toc_page_map else ""
            toc_rows.append([
                Paragraph(title_text, styles['TOCLeft']),
                Paragraph(f"<b>{page_str}</b>", styles['TOCRight']),
            ])
        if toc_rows:
            t = Table(toc_rows, colWidths=[content_w - 0.6*inch, 0.6*inch])
            t.setStyle(TableStyle([
                ('VALIGN',(0,0),(-1,-1),'TOP'),
                ('TOPPADDING',(0,0),(-1,-1),4),
                ('BOTTOMPADDING',(0,0),(-1,-1),4),
                ('LINEBELOW',(0,0),(-1,-1),0.3,colors.Color(.85,.85,.85)),
            ]))
            story.append(t)
        story.append(PageBreak())
        
        # Chapters
        for ch in chapters:
            cn = ch['chapter_number']
            # Marker (invisible, records page number)
            story.append(ChapterMark(cn, page_tracker))
            # Chapter title page
            story.append(Spacer(1, 2.5 * inch))
            lbl = f"CHAPITRE {cn}" if is_fr else f"CHAPTER {cn}"
            story.append(Paragraph(lbl, styles['ChapLabel']))
            story.append(Spacer(1, 12))
            story.append(Paragraph(md_to_xml(ch['title']), styles['ChapTitlePage']))
            story.append(PageBreak())
            # Chapter body
            story.extend(build_chapter_body(ch, styles))
            story.append(PageBreak())
        return story
    
    # ---- Page number drawing ----
    def draw_page_number(canvas, doc):
        pn = canvas.getPageNumber()
        if pn > 1:
            canvas.saveState()
            canvas.setFont("Helvetica", 9)
            canvas.setFillColor(colors.Color(0.4, 0.4, 0.4))
            canvas.drawCentredString(page_w / 2, 0.4 * inch, str(pn))
            canvas.restoreState()
    
    # ===== PASS 1: Build to get real page numbers =====
    page_tracker = {}
    styles1 = make_styles()
    story1 = build_story(styles1, toc_page_map=None)
    
    buf1 = BytesIO()
    doc1 = BaseDocTemplate(buf1, pagesize=(page_w, page_h),
        leftMargin=left_m, rightMargin=right_m, topMargin=top_m, bottomMargin=bottom_m)
    frame1 = Frame(left_m, bottom_m, content_w, page_h - top_m - bottom_m, id='main')
    doc1.addPageTemplates([
        PageTemplate(id='first', frames=[frame1], onPage=lambda c,d: None),
        PageTemplate(id='later', frames=[frame1], onPage=draw_page_number),
    ])
    # Switch to 'later' template after first page
    story1.insert(1, Spacer(0, 0))  # ensure the template switch happens
    
    try:
        doc1.build(story1)
    except Exception as e:
        logger.error(f"PDF pass 1 error: {e}")
    
    logger.info(f"PDF Pass 1 - Chapter page map: {page_tracker}")
    
    # ===== PASS 2: Rebuild with correct TOC page numbers =====
    page_tracker_pass2 = {}
    # Re-point ChapterMark to new dict so pass2 markers don't overwrite
    # We'll use the pass1 results
    final_page_map = dict(page_tracker)
    
    # Build a new tracker for pass 2 (won't be used for TOC but keeps markers happy)
    page_tracker.clear()
    
    styles2 = make_styles()
    story2 = build_story(styles2, toc_page_map=final_page_map)
    
    doc2 = BaseDocTemplate(str(filepath), pagesize=(page_w, page_h),
        leftMargin=left_m, rightMargin=right_m, topMargin=top_m, bottomMargin=bottom_m)
    frame2 = Frame(left_m, bottom_m, content_w, page_h - top_m - bottom_m, id='main')
    doc2.addPageTemplates([
        PageTemplate(id='first', frames=[frame2], onPage=lambda c,d: None),
        PageTemplate(id='later', frames=[frame2], onPage=draw_page_number),
    ])
    
    doc2.build(story2)
    return filepath

async def export_docx(book):
    """Generate KDP-compliant DOCX with proper formatting, chapter title pages, TOC with page numbers."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    filepath = EXPORTS_DIR / f"{book['id']}.docx"
    doc = Document()
    is_fr = book.get('language') != 'en'
    
    # Page setup - KDP 5.5 x 8.5
    section = doc.sections[0]
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    
    # Footer with page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin'); run._r.append(fld1)
    run2 = footer_para.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '; run2._r.append(instr)
    run3 = footer_para.add_run()
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end'); run3._r.append(fld2)
    
    chapters = sorted(book.get('chapters', []), key=lambda x: x.get('chapter_number', 0))
    
    # ---- PRE-CALCULATE real page numbers ----
    # KDP 5.5x8.5 with margins => usable area ~4.25 x 7 inches
    # At 11pt Times Roman ~16pt leading, about 31 lines per page, ~10 words/line => ~310 words/page
    # A heading takes ~2 lines, a list item ~1.5 lines, image ~12 lines
    WORDS_PER_PAGE = 280  # conservative for formatted text
    LINES_PER_PAGE = 31
    
    chapter_page_starts = {}
    current_page = 3  # page 1 = title, page 2 = TOC (at least), page 3+ = chapters
    
    # Estimate TOC pages: ~1.5 lines per entry
    toc_lines = len(chapters) * 1.5 + 4  # header + spacing
    toc_pages = max(1, int(toc_lines / LINES_PER_PAGE) + 1)
    current_page = 1 + toc_pages + 1  # title page + toc pages + 1
    
    for ch in chapters:
        chapter_page_starts[ch['chapter_number']] = current_page
        current_page += 1  # chapter title page
        
        content = strip_chapter_title_from_content(ch.get('content', ''), ch.get('title', ''))
        lines = content.split('\n')
        
        # Count estimated lines
        total_lines = 0
        if ch.get('image_url'):
            total_lines += 14  # image takes ~14 lines
        for line in lines:
            lt, lc, lv = parse_markdown_line(line)
            if lt == "blank":
                total_lines += 0.5
            elif lt == "heading":
                total_lines += 3
            elif lt in ("list_item", "num_list_item"):
                total_lines += 1.2
            elif lt == "paragraph":
                word_count = len(lc.split())
                total_lines += max(1, word_count / 10)
            elif lt == "hr":
                total_lines += 2
        
        content_pages = max(1, int(total_lines / LINES_PER_PAGE) + 1)
        current_page += content_pages
    
    # ---- TITLE PAGE ----
    for _ in range(8):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(md_clean(book['title']))
    run.font.size = Pt(28); run.bold = True; run.font.name = 'Georgia'
    if book.get('subtitle'):
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run(md_clean(book['subtitle']))
        run.font.size = Pt(14); run.font.color.rgb = RGBColor(100,100,100); run.font.name = 'Georgia'
    doc.add_page_break()
    
    # ---- TOC with 2-column table ----
    toc_title = "Table des matieres" if is_fr else "Table of Contents"
    h = doc.add_heading(toc_title, level=1)
    for r in h.runs: r.font.name = 'Georgia'
    
    toc_table = doc.add_table(rows=len(chapters), cols=2)
    toc_table.allow_autofit = True
    
    for row_idx, ch in enumerate(chapters):
        ch_lbl = f"Chapitre {ch['chapter_number']}" if is_fr else f"Chapter {ch['chapter_number']}"
        page_num = chapter_page_starts.get(ch['chapter_number'], "")
        
        left_cell = toc_table.cell(row_idx, 0)
        left_cell.text = ""
        lp = left_cell.paragraphs[0]
        run = lp.add_run(f"{ch_lbl}  -  {md_clean(ch['title'])}")
        run.font.name = 'Georgia'; run.font.size = Pt(11)
        
        right_cell = toc_table.cell(row_idx, 1)
        right_cell.text = ""
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = rp.add_run(str(page_num))
        run.font.name = 'Georgia'; run.font.size = Pt(11); run.bold = True
    
    # Remove table borders
    tbl = toc_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0'); borders.append(b)
    tblPr.append(borders)
    
    doc.add_page_break()
    
    # ---- CHAPTERS ----
    for chapter in chapters:
        cn = chapter['chapter_number']
        ch_lbl = f"Chapitre {cn}" if is_fr else f"Chapter {cn}"
        
        # Chapter title page
        for _ in range(8):
            doc.add_paragraph()
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(ch_lbl.upper())
        lr.font.size = Pt(10); lr.font.color.rgb = RGBColor(128,128,128); lr.font.name = 'Arial'
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(md_clean(chapter['title']))
        tr.font.size = Pt(22); tr.bold = True; tr.font.name = 'Georgia'
        doc.add_page_break()
        
        # Chapter image
        if chapter.get('image_url') and chapter['image_url'].startswith('/api/images/'):
            img_filename = chapter['image_url'].replace('/api/images/', '')
            img_path = IMAGES_DIR / img_filename
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(3.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception: pass
        
        # Chapter content (stripped)
        content = strip_chapter_title_from_content(chapter.get('content', ''), chapter.get('title', ''))
        for line in content.split('\n'):
            lt, lc, lv = parse_markdown_line(line)
            cleaned = md_clean(lc)
            if lt == "blank": continue
            elif lt == "heading":
                h = doc.add_heading(cleaned, level=min(lv+1, 4))
                for r in h.runs: r.font.name = 'Georgia'
            elif lt in ("list_item", "num_list_item"):
                p = doc.add_paragraph(style='List Bullet' if lt == "list_item" else 'List Number')
                _add_formatted_runs(p, lc); p.paragraph_format.space_after = Pt(3)
            elif lt == "hr":
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("_" * 30); r.font.color.rgb = RGBColor(180,180,180)
            elif lt == "paragraph":
                p = doc.add_paragraph()
                _add_formatted_runs(p, lc); p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()
    
    doc.save(str(filepath))
    return filepath

def _add_formatted_runs(paragraph, text):
    """Add runs with bold/italic formatting from markdown to a DOCX paragraph."""
    from docx.shared import Pt
    
    # Process bold+italic, bold, italic patterns
    parts = re.split(r'(\*{2,3}.+?\*{2,3}|_{2,3}.+?_{2,3})', text)
    for part in parts:
        if re.match(r'^\*{3}(.+?)\*{3}$', part) or re.match(r'^_{3}(.+?)_{3}$', part):
            clean = re.sub(r'^[*_]{3}|[*_]{3}$', '', part)
            run = paragraph.add_run(clean)
            run.bold = True
            run.italic = True
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
        elif re.match(r'^\*{2}(.+?)\*{2}$', part) or re.match(r'^_{2}(.+?)_{2}$', part):
            clean = re.sub(r'^[*_]{2}|[*_]{2}$', '', part)
            run = paragraph.add_run(clean)
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
        else:
            # Could still contain single * italic
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sp in sub_parts:
                if re.match(r'^\*([^*]+?)\*$', sp):
                    clean = sp.strip('*')
                    run = paragraph.add_run(clean)
                    run.italic = True
                    run.font.name = 'Georgia'
                    run.font.size = Pt(11)
                else:
                    # Strip remaining markdown
                    cleaned = md_clean(sp)
                    if cleaned:
                        run = paragraph.add_run(cleaned)
                        run.font.name = 'Georgia'
                        run.font.size = Pt(11)

async def export_epub(book):
    """Generate EPUB with proper formatting, chapter title pages, TOC."""
    from ebooklib import epub
    
    filepath = EXPORTS_DIR / f"{book['id']}.epub"
    is_fr = book.get('language') != 'en'
    
    ebook = epub.EpubBook()
    ebook.set_identifier(book['id'])
    ebook.set_title(book['title'])
    ebook.set_language(book.get('language', 'fr'))
    
    css_content = """
    body { font-family: Georgia, 'Times New Roman', serif; line-height: 1.8; color: #1a1a1a; margin: 1em; }
    h1 { font-size: 1.6em; margin-top: 2em; margin-bottom: 0.5em; font-weight: bold; }
    h2 { font-size: 1.3em; margin-top: 1.5em; margin-bottom: 0.4em; font-weight: bold; }
    h3 { font-size: 1.1em; margin-top: 1.2em; margin-bottom: 0.3em; font-weight: bold; font-style: italic; }
    h4 { font-size: 1em; margin-top: 1em; font-weight: bold; }
    p { text-align: justify; margin-bottom: 0.5em; font-size: 1em; }
    ul, ol { margin-left: 1.5em; margin-bottom: 0.5em; }
    li { margin-bottom: 0.2em; }
    .chapter-title-page { text-align: center; padding-top: 40%; }
    .chapter-label { font-size: 0.85em; color: #888; letter-spacing: 0.2em; text-transform: uppercase; margin-bottom: 0.5em; }
    .chapter-title { font-size: 1.8em; margin-top: 0.2em; font-weight: bold; }
    hr { border: none; border-top: 1px solid #ccc; margin: 1.5em 20%; }
    strong { font-weight: bold; }
    em { font-style: italic; }
    code { font-family: 'Courier New', monospace; font-size: 0.9em; background: #f5f5f5; padding: 0.1em 0.3em; }
    .toc-table { width: 100%; border-collapse: collapse; }
    .toc-table td { padding: 0.3em 0; vertical-align: top; }
    .toc-table td:last-child { text-align: right; font-weight: bold; width: 3em; }
    """
    style = epub.EpubItem(uid="style", file_name="style/default.css", 
                          media_type="text/css", content=css_content.encode('utf-8'))
    ebook.add_item(style)
    
    chapters_epub = []
    chapters = sorted(book.get('chapters', []), key=lambda x: x.get('chapter_number', 0))
    
    # TOC page with table layout
    toc_ch = epub.EpubHtml(title="Table des matieres" if is_fr else "Table of Contents",
                           file_name="toc.xhtml", lang=book.get('language', 'fr'))
    toc_label = "Table des matieres" if is_fr else "Table of Contents"
    toc_html = f"<h1>{toc_label}</h1><table class='toc-table'>"
    for idx, ch_data in enumerate(chapters):
        ch_lbl = f"Chapitre {ch_data['chapter_number']}" if is_fr else f"Chapter {ch_data['chapter_number']}"
        toc_html += f'<tr><td><a href="chapter_{ch_data["chapter_number"]}.xhtml">{ch_lbl}  -  {md_to_html(ch_data["title"])}</a></td><td>{idx + 3}</td></tr>'
    toc_html += "</table>"
    toc_ch.content = toc_html
    toc_ch.add_item(style)
    ebook.add_item(toc_ch)
    
    for chapter in chapters:
        ch = epub.EpubHtml(
            title=chapter['title'],
            file_name=f"chapter_{chapter['chapter_number']}.xhtml",
            lang=book.get('language', 'fr')
        )
        
        ch_label = f"Chapitre {chapter['chapter_number']}" if is_fr else f"Chapter {chapter['chapter_number']}"
        
        # Dedicated chapter title section
        content_html = f'<div class="chapter-title-page">'
        content_html += f'<p class="chapter-label">{ch_label}</p>'
        content_html += f'<h1 class="chapter-title">{md_to_html(chapter["title"])}</h1>'
        content_html += f'</div><hr/>'
        
        # Content with title stripped
        stripped_content = strip_chapter_title_from_content(
            chapter.get('content', ''), chapter.get('title', ''))
        
        in_list = False
        list_type = None
        
        for line in stripped_content.split('\n'):
            line_type, line_content, level = parse_markdown_line(line)
            
            if line_type in ("list_item", "num_list_item"):
                new_list_type = "ul" if line_type == "list_item" else "ol"
                if not in_list or list_type != new_list_type:
                    if in_list:
                        content_html += f"</{list_type}>"
                    content_html += f"<{new_list_type}>"
                    in_list = True
                    list_type = new_list_type
                content_html += f"<li>{md_to_html(line_content)}</li>"
                continue
            
            if in_list:
                content_html += f"</{list_type}>"
                in_list = False
                list_type = None
            
            if line_type == "blank":
                continue
            elif line_type == "heading":
                tag = f"h{min(level + 1, 4)}"
                content_html += f"<{tag}>{md_to_html(line_content)}</{tag}>"
            elif line_type == "hr":
                content_html += "<hr/>"
            elif line_type == "paragraph":
                content_html += f"<p>{md_to_html(line_content)}</p>"
        
        if in_list:
            content_html += f"</{list_type}>"
        
        ch.content = content_html
        ch.add_item(style)
        ebook.add_item(ch)
        chapters_epub.append(ch)
    
    ebook.toc = [toc_ch] + [(epub.Section(ch.title), [ch]) for ch in chapters_epub]
    ebook.add_item(epub.EpubNcx())
    ebook.add_item(epub.EpubNav())
    ebook.spine = ['nav', toc_ch] + chapters_epub
    
    epub.write_epub(str(filepath), ebook)
    return filepath

# ====== ROOT ======

@api_router.get("/")
async def root():
    return {"message": "Lumina Press API", "version": "1.0.0"}

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
